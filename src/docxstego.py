import os
from docx import Document


def hide_text(cover_file, payload):
    doc_cover = Document(cover_file)

    # Convert payload to binary
    payload_bin = ''.join(format(ord(c), '08b') for c in payload)

    # Define non-printing whitespace characters for 0, 1, and boundary separator
    zero_char = '\u200B'  # Zero-width space
    one_char = '\u200C'  # Zero-width non-joiner
    boundary_char = '\u200D'  # Zero-width joiner

    # Generate the stego text by replacing each bit with corresponding characters
    stego_text = boundary_char
    for bit in payload_bin:
        if bit == '0':
            stego_text += zero_char
        else:
            stego_text += one_char
        # stego_text += boundary_char

    # Concatenate the cover text and stego text
    doc_cover_text = "\n".join([paragraph.text for paragraph in doc_cover.paragraphs])
    stego_text = doc_cover_text + stego_text

    # Create a new document for the stego file
    stego_doc = Document()
    stego_doc.add_paragraph(stego_text)

    root, ext = os.path.splitext(cover_file)
    directory_path = os.path.dirname(root)
    basename_without_ext = os.path.splitext(os.path.basename(cover_file))[0]
    outfileName = basename_without_ext + "_encoded" + ext
    stego_file = os.path.join(directory_path, outfileName)
    stego_doc.save(stego_file)

    print("SUCCESS MESSAGE: Successful payload hidden in the document file.")

def extract_text(stego_file):
    doc_stego = Document(stego_file)

    # Extract the stego text from the document
    stego_text = "\n".join([paragraph.text for paragraph in doc_stego.paragraphs])

    # Find the start index of the payload
    start_index = stego_text.find('\u200D') + 1

    # Extract the payload from the stego text
    payload_text = stego_text[start_index:]

    # Convert payload from characters to binary
    payload_bin = ''
    for char in payload_text:
        if char == '\u200B':
            payload_bin += '0'
        elif char == '\u200C':
            payload_bin += '1'
        else:
            payload_bin += ' '

    # Convert payload from binary to text
    extracted_text = ''
    for i in range(0, len(payload_bin), 8):
        byte = payload_bin[i:i + 8]
        extracted_text += chr(int(byte, 2))

    return extracted_text

if __name__ == '__main__':
    payload = "This is a hidden message."
    cover_file = "Docx/doc_cover.docx"

    hide_text(cover_file, payload)

    stego_file = "Docx/doc_cover_encoded.docx"
    extracted_text = extract_text(stego_file)
    print("EXTRACTED MESSAGE:", extracted_text)
