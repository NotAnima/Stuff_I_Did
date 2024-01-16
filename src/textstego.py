import os
from docx import Document

def hide_text(cover_file, payload, ext):
    if ext == ".docx":
        doc_cover = Document(cover_file)
    else:
        with open(cover_file, 'r') as cover:
            cover_text = cover.read()

    # Convert payload to binary
    payload_bin = ''.join(format(ord(c), '08b') for c in payload)

    # Define non-printing whitespace characters for 0, 1, and boundary separator
    zero_char = '\u200B'  # Zero-width space
    one_char = '\u200C'  # Zero-width non-joiner
    boundary_char = '\u200D'  # Zero-width joiner

    # Generate the stego text by replacing each bit with corresponding characters
    stego_text = boundary_char
    for bit in payload_bin:
        if bit == '0':
            stego_text += zero_char
        else:
            stego_text += one_char
        # stego_text += boundary_char

    root, ext = os.path.splitext(cover_file)
    directory_path = os.path.dirname(root)
    basename_without_ext = os.path.splitext(os.path.basename(cover_file))[0]
    outfileName = basename_without_ext + "_encoded" + ext
    stego_file = os.path.join(directory_path, outfileName)

    # Concatenate the cover text and stego text
    if ext == ".docx":
        doc_cover_text = "\n".join([paragraph.text for paragraph in doc_cover.paragraphs])
        stego_doc = Document()
        stego_doc.add_paragraph(stego_text)
        stego_doc.save(stego_file)
    else:
        stego_text = cover_text + stego_text
        with open(stego_file, 'w', encoding='utf-8') as stego:
            stego.write(stego_text)

    print("Text payload hidden successfully in the cover file.")
    return stego_file

def extract_text(stego_file, ext):
    if ext == ".docx":
        doc_stego = Document(stego_file)
        stego_text = "\n".join([paragraph.text for paragraph in doc_stego.paragraphs])
    else:
        with open(stego_file, 'r', encoding='utf-8') as stego:
            stego_text = stego.read()

    # Find the start index of the payload
    start_index = stego_text.find('\u200D') + 1

    # Extract the payload from the stego text
    payload_chars = stego_text[start_index:]

    # Convert payload from characters to binary
    payload_bin = ''
    for char in payload_chars:
        if char == '\u200B':
            payload_bin += '0'
        elif char == '\u200C':
            payload_bin += '1'
        else:
            payload_bin += ' '

    # Convert payload from binary to text
    extracted_text = ''
    for i in range(0, len(payload_bin), 8):
        byte = payload_bin[i:i + 8]
        extracted_text += chr(int(byte, 2))

    return extracted_text

# # Example usage
# payload = "This is a hidden message."
# cover_file = "Text/cover.txt"
#
# # Hide the text payload in the cover file
# hide_text(cover_file, payload)
#
# stego_file = "cover_encoded.txt"
# # Extract the hidden text from the stego file
# extracted_text = extract_text(stego_file)
# print("Extracted text:", extracted_text)
